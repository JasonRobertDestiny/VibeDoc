"""
VibeDoc 多格式导出管理器
支持 Ma# PDF 导出
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# 高级PDF导出 - 移除weasyprint依赖，使用reportlab
WEASYPRINT_AVAILABLE = FalseF 格式的文档导出
"""

import os
import io
import re
import zipfile
import tempfile
from datetime import datetime
from typing import Dict, Tuple, Optional, Any
import logging

# 核心依赖
import markdown
import html2text

# Word 导出
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF 导出
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# 高级PDF导出（备用方案） - 移除weasyprint依赖
WEASYPRINT_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportManager:
    """多格式导出管理器"""
    
    def __init__(self):
        self.supported_formats = ['markdown', 'html']
        
        # 检查可选依赖
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')
        if PDF_AVAILABLE:
            self.supported_formats.append('pdf')
            
        logger.info(f"📄 ExportManager 初始化完成，支持格式: {', '.join(self.supported_formats)}")
    
    def get_supported_formats(self) -> list:
        """获取支持的导出格式"""
        return self.supported_formats.copy()
    
    def export_to_markdown(self, content: str, metadata: Optional[Dict] = None) -> str:
        """
        导出为 Markdown 格式（清理和优化）
        
        Args:
            content: 原始内容
            metadata: 元数据信息
            
        Returns:
            str: 优化后的 Markdown 内容
        """
        try:
            # 添加文档头部信息
            if metadata:
                header = f"""---
title: {metadata.get('title', 'VibeDoc 开发计划')}
author: {metadata.get('author', 'VibeDoc AI Agent')}
date: {metadata.get('date', datetime.now().strftime('%Y-%m-%d'))}
generator: VibeDoc AI Agent v1.0
---

"""
                content = header + content
            
            # 清理和优化内容
            content = self._clean_markdown_content(content)
            
            logger.info("✅ Markdown 导出成功")
            return content
            
        except Exception as e:
            logger.error(f"❌ Markdown 导出失败: {e}")
            return content  # 返回原始内容
    
    def export_to_html(self, content: str, metadata: Optional[Dict] = None) -> str:
        """
        导出为 HTML 格式（带样式）
        
        Args:
            content: Markdown 内容
            metadata: 元数据信息
            
        Returns:
            str: 完整的 HTML 内容
        """
        try:
            # 配置 Markdown 扩展
            md = markdown.Markdown(
                extensions=[
                    'markdown.extensions.extra',
                    'markdown.extensions.codehilite',
                    'markdown.extensions.toc',
                    'markdown.extensions.tables'
                ],
                extension_configs={
                    'codehilite': {
                        'css_class': 'highlight',
                        'use_pygments': False
                    },
                    'toc': {
                        'title': '目录'
                    }
                }
            )
            
            # 转换 Markdown 到 HTML
            html_content = md.convert(content)
            
            # 生成完整的 HTML 文档
            title = metadata.get('title', 'VibeDoc 开发计划') if metadata else 'VibeDoc 开发计划'
            author = metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'
            date = metadata.get('date', datetime.now().strftime('%Y-%m-%d')) if metadata else datetime.now().strftime('%Y-%m-%d')
            
            full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <meta name="author" content="{author}">
    <meta name="generator" content="VibeDoc AI Agent">
    <style>
        {self._get_html_styles()}
    </style>
    <!-- Mermaid 支持 -->
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10.6.1/dist/mermaid.min.js"></script>
    <script>
        document.addEventListener('DOMContentLoaded', function() {{
            mermaid.initialize({{ 
                startOnLoad: true,
                theme: 'default',
                securityLevel: 'loose',
                flowchart: {{ useMaxWidth: true }}
            }});
        }});
    </script>
</head>
<body>
    <div class="container">
        <header class="document-header">
            <h1>{title}</h1>
            <div class="meta-info">
                <span class="author">📝 {author}</span>
                <span class="date">📅 {date}</span>
                <span class="generator">🤖 Generated by VibeDoc AI Agent</span>
            </div>
        </header>
        
        <main class="content">
            {html_content}
        </main>
        
        <footer class="document-footer">
            <p>本文档由 <strong>VibeDoc AI Agent</strong> 生成 | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </footer>
    </div>
</body>
</html>"""
            
            logger.info("✅ HTML 导出成功")
            return full_html
            
        except Exception as e:
            logger.error(f"❌ HTML 导出失败: {e}")
            # 简单的 HTML 备用方案
            return f"""<!DOCTYPE html>
<html><head><title>VibeDoc 开发计划</title></head>
<body><pre>{content}</pre></body></html>"""
    
    def export_to_docx(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """
        导出为 Word 文档格式
        
        Args:
            content: Markdown 内容
            metadata: 元数据信息
            
        Returns:
            bytes: Word 文档的二进制数据
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，无法导出 Word 格式")
        
        try:
            # 创建新文档
            doc = Document()
            
            # 设置文档属性
            properties = doc.core_properties
            properties.title = metadata.get('title', 'VibeDoc 开发计划') if metadata else 'VibeDoc 开发计划'
            properties.author = metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'
            properties.subject = 'AI驱动的智能开发计划'
            properties.comments = 'Generated by VibeDoc AI Agent'
            
            # 添加标题
            title = doc.add_heading(properties.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元信息
            doc.add_paragraph()
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"📝 作者: {properties.author}").bold = True
            meta_para.add_run("\n")
            meta_para.add_run(f"📅 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
            meta_para.add_run("\n")
            meta_para.add_run("🤖 生成工具: VibeDoc AI Agent").bold = True
            
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
            
            # 解析和添加内容
            self._parse_markdown_to_docx(doc, content)
            
            # 添加页脚
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            footer_para = doc.add_paragraph()
            footer_para.add_run("本文档由 VibeDoc AI Agent 自动生成").italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存到内存
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            
            logger.info("✅ Word 文档导出成功")
            return doc_stream.getvalue()
            
        except Exception as e:
            logger.error(f"❌ Word 导出失败: {e}")
            raise
    
    def export_to_pdf(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """
        导出为 PDF 格式
        
        Args:
            content: Markdown 内容  
            metadata: 元数据信息
            
        Returns:
            bytes: PDF 文档的二进制数据
        """
        if PDF_AVAILABLE:
            return self._export_pdf_reportlab(content, metadata)
        else:
            raise ImportError("PDF 导出依赖未安装")
    
    def create_multi_format_export(self, content: str, formats: list = None, metadata: Optional[Dict] = None) -> bytes:
        """
        创建多格式导出的 ZIP 包
        
        Args:
            content: 原始内容
            formats: 要导出的格式列表，默认为所有支持的格式
            metadata: 元数据信息
            
        Returns:
            bytes: ZIP 文件的二进制数据
        """
        if formats is None:
            formats = self.supported_formats
        
        # 验证格式
        invalid_formats = set(formats) - set(self.supported_formats)
        if invalid_formats:
            raise ValueError(f"不支持的格式: {', '.join(invalid_formats)}")
        
        try:
            # 创建内存中的 ZIP 文件
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # 生成基础文件名
                base_name = metadata.get('title', 'vibedoc_plan') if metadata else 'vibedoc_plan'
                base_name = re.sub(r'[^\w\-_\.]', '_', base_name)  # 清理文件名
                
                # 导出各种格式
                for fmt in formats:
                    try:
                        if fmt == 'markdown':
                            file_content = self.export_to_markdown(content, metadata)
                            zip_file.writestr(f"{base_name}.md", file_content.encode('utf-8'))
                            
                        elif fmt == 'html':
                            file_content = self.export_to_html(content, metadata)
                            zip_file.writestr(f"{base_name}.html", file_content.encode('utf-8'))
                            
                        elif fmt == 'docx' and DOCX_AVAILABLE:
                            file_content = self.export_to_docx(content, metadata)
                            zip_file.writestr(f"{base_name}.docx", file_content)
                            
                        elif fmt == 'pdf' and PDF_AVAILABLE:
                            file_content = self.export_to_pdf(content, metadata)
                            zip_file.writestr(f"{base_name}.pdf", file_content)
                            
                    except Exception as e:
                        logger.warning(f"⚠️ 格式 {fmt} 导出失败: {e}")
                        # 在 ZIP 中添加错误信息文件
                        error_msg = f"格式 {fmt} 导出失败:\n{str(e)}\n\n请检查相关依赖是否正确安装。"
                        zip_file.writestr(f"ERROR_{fmt}.txt", error_msg.encode('utf-8'))
                
                # 添加说明文件
                readme_content = f"""# VibeDoc 导出文件包

## 📋 文件说明
本压缩包包含了您的开发计划的多种格式导出：

### 📄 支持的格式：
- **Markdown (.md)**: 原始格式，支持所有 Markdown 语法
- **HTML (.html)**: 网页格式，包含样式和 Mermaid 图表支持
- **Word (.docx)**: Microsoft Word 文档格式
- **PDF (.pdf)**: 便携式文档格式

### 🤖 生成信息：
- 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- 生成工具: VibeDoc AI Agent v1.0
- 项目地址: https://github.com/JasonRobertDestiny/VibeDocs

### 💡 使用建议：
1. 优先使用 HTML 格式查看，支持最佳的视觉效果
2. 使用 Markdown 格式进行进一步编辑
3. 使用 Word 格式进行正式文档处理
4. 使用 PDF 格式进行分享和打印

---
感谢使用 VibeDoc AI Agent！
"""
                zip_file.writestr("README.md", readme_content.encode('utf-8'))
            
            zip_buffer.seek(0)
            logger.info(f"✅ 多格式导出成功，包含 {len(formats)} 种格式")
            return zip_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"❌ 多格式导出失败: {e}")
            raise
    
    def _clean_markdown_content(self, content: str) -> str:
        """清理和优化 Markdown 内容"""
        # 修复常见的格式问题
        content = re.sub(r'\n{3,}', '\n\n', content)  # 移除多余空行
        content = re.sub(r'(?m)^[ \t]+$', '', content)  # 移除只有空格的行
        content = content.strip()
        
        return content
    
    def _get_html_styles(self) -> str:
        """获取 HTML 样式"""
        return """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f8fafc;
        }
        
        .container {
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            border-radius: 8px;
            margin-top: 20px;
            margin-bottom: 20px;
        }
        
        .document-header {
            text-align: center;
            border-bottom: 3px solid #667eea;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        
        .document-header h1 {
            color: #667eea;
            font-size: 2.2em;
            margin-bottom: 15px;
        }
        
        .meta-info {
            display: flex;
            justify-content: center;
            gap: 20px;
            flex-wrap: wrap;
            color: #666;
            font-size: 0.9em;
        }
        
        .content h1, .content h2, .content h3, .content h4 {
            color: #2d3748;
            margin-top: 2em;
            margin-bottom: 1em;
        }
        
        .content h1 { border-bottom: 2px solid #667eea; padding-bottom: 0.5em; }
        .content h2 { border-bottom: 1px solid #e2e8f0; padding-bottom: 0.3em; }
        
        .content p {
            margin-bottom: 1em;
            text-align: justify;
        }
        
        .content ul, .content ol {
            margin-bottom: 1em;
            padding-left: 2em;
        }
        
        .content li {
            margin-bottom: 0.5em;
        }
        
        .content pre {
            background: #2d3748;
            color: #e2e8f0;
            padding: 1em;
            border-radius: 6px;
            overflow-x: auto;
            margin: 1em 0;
        }
        
        .content code {
            background: #f7fafc;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: 'SFMono-Regular', Consolas, monospace;
        }
        
        .content table {
            width: 100%;
            border-collapse: collapse;
            margin: 1em 0;
        }
        
        .content th, .content td {
            border: 1px solid #e2e8f0;
            padding: 0.75em;
            text-align: left;
        }
        
        .content th {
            background: #f7fafc;
            font-weight: 600;
        }
        
        .content blockquote {
            border-left: 4px solid #667eea;
            margin: 1em 0;
            padding-left: 1em;
            color: #666;
            font-style: italic;
        }
        
        .mermaid {
            text-align: center;
            margin: 2em 0;
        }
        
        .document-footer {
            margin-top: 3em;
            padding-top: 20px;
            border-top: 1px solid #e2e8f0;
            text-align: center;
            color: #666;
            font-size: 0.9em;
        }
        
        @media (max-width: 768px) {
            .container {
                margin: 10px;
                padding: 15px;
            }
            
            .meta-info {
                flex-direction: column;
                gap: 10px;
            }
        }
        """
    
    def _parse_markdown_to_docx(self, doc: "Document", content: str):
        """解析 Markdown 内容并添加到 Word 文档"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
                
            # 标题处理
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                if level <= 6:
                    doc.add_heading(title_text, level)
                    continue
            
            # 代码块处理（简化）
            if line.startswith('```'):
                continue
                
            # 列表处理
            if line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                para = doc.add_paragraph(text, style='List Bullet')
                continue
                
            if re.match(r'^\d+\.', line):
                text = re.sub(r'^\d+\.\s*', '', line)
                para = doc.add_paragraph(text, style='List Number')
                continue
            
            # 普通段落
            if line:
                # 简单的粗体和斜体处理
                line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # 移除粗体标记，Word 中后续可以手动设置
                line = re.sub(r'\*(.*?)\*', r'\1', line)      # 移除斜体标记
                doc.add_paragraph(line)
    
    def _export_pdf_reportlab(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """使用 ReportLab 导出 PDF"""
        try:
            buffer = io.BytesIO()
            
            # 创建 PDF 文档
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                topMargin=1*inch,
                bottomMargin=1*inch,
                leftMargin=1*inch,
                rightMargin=1*inch
            )
            
            # 样式设置
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=20,
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            # 构建内容
            story = []
            
            # 添加标题
            title = metadata.get('title', 'VibeDoc 开发计划') if metadata else 'VibeDoc 开发计划'
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # 添加元信息
            meta_text = f"""
            作者: {metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'}<br/>
            生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
            生成工具: VibeDoc AI Agent
            """
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 30))
            
            # 简单处理 Markdown 内容（基础版本）
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                    continue
                    
                if line.startswith('#'):
                    # 标题
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    if level == 1:
                        story.append(Paragraph(title_text, styles['Heading1']))
                    elif level == 2:
                        story.append(Paragraph(title_text, styles['Heading2']))
                    else:
                        story.append(Paragraph(title_text, styles['Heading3']))
                else:
                    # 普通段落
                    story.append(Paragraph(line, styles['Normal']))
                    
                story.append(Spacer(1, 6))
            
            # 生成 PDF
            doc.build(story)
            buffer.seek(0)
            
            logger.info("✅ PDF 导出成功（ReportLab）")
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"❌ ReportLab PDF 导出失败: {e}")
            raise

# 全局导出管理器实例
export_manager = ExportManager()